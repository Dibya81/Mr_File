from docx import Document


class WordExtractor:
    def extract(self, file_path: str) -> dict:
        try:
            doc = Document(file_path)

            paragraphs = len(doc.paragraphs)
            text = "\n".join([p.text for p in doc.paragraphs])
            word_count = len(text.split()) if text.strip() else 0
            char_count = len(text)
            table_count = len(doc.tables)

            core_props = doc.core_properties
            result = {
                "paragraph_count": paragraphs,
                "word_count": word_count,
                "character_count": char_count,
                "table_count": table_count,
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "text_preview": text[:500] if text else "",
            }
            return result
        except Exception as e:
            return {"error": str(e), "paragraph_count": 0, "word_count": 0}


word_extractor = WordExtractor()
